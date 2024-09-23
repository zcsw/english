# This Python file is used to format or optimize the layout of tables for better readability and space efficiency.
# Author: POE@Claude-3.5-Sonnet

import os
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def process_docx_tables(input_file):
    # Backup original file
    backup_file = input_file + '.backup'
    os.rename(input_file, backup_file)

    # Open the document
    doc = Document(backup_file)

    # Process each table in the document
    for table in doc.tables:
        # Remove inner borders, keep outer borders
        tbl_pr = table._element.xpath('w:tblPr')[0]
        tbl_borders = parse_xml(f'<w:tblBorders {nsdecls("w")}>'
                                '<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                '<w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                '<w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                '<w:insideH w:val="none"/>'
                                '<w:insideV w:val="none"/>'
                                '</w:tblBorders>')
        tbl_pr.append(tbl_borders)

        # Process each cell in the table
        for row in table.rows:
            for cell in row.cells:
                # Set vertical alignment to top
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                # Set spacing before and after to 0 for each paragraph in the cell
                for paragraph in cell.paragraphs:
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_before = Pt(0)
                    paragraph_format.space_after = Pt(0)
                    # Ensure the spacing is applied
                    paragraph_format.line_spacing = 1.0
                    # Clear any existing spacing
                    if hasattr(paragraph_format, '_element'):
                        spacing_element = paragraph_format._element.xpath('./w:spacing')
                        if spacing_element:
                            paragraph_format._element.remove(spacing_element[0])

    # Save the modified document with the original filename
    doc.save(input_file)
    print(f"Document processed and saved as {input_file}")
    print(f"Original file backed up as {backup_file}")

# Usage
input_file = "2016考研英语一真题.docx"  # Replace with your input file name
process_docx_tables(input_file)


