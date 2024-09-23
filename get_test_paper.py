# This Python file aims to generate test papers by removing Chinese text from tables in Word documents.
# Author: POE@Claude-3.5-Sonnet
# There is a bug where a table remains after Part B because the code only deletes the text and not the table.

import os
from docx import Document
from docx2pdf import convert
import re
from PyPDF2 import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from io import BytesIO

pdfmetrics.registerFont(TTFont('SimSun', 'simsun.ttc'))

def process_docx(input_file):
    # Create a new Document object without modifying the original file
    doc = Document(input_file)
    
    # Delete "Part B" and everything after it
    delete_from = None
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith("Part B"):
            delete_from = i
            break
    
    if delete_from is not None:
        for i in range(len(doc.paragraphs) - 1, delete_from - 1, -1):
            delete_paragraph(doc.paragraphs[i])
    
    # Clear the second column of all tables
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                row.cells[1].text = ""
    
    # Save as a new docx file
    base_name = os.path.splitext(os.path.basename(input_file))[0]
    temp_docx = f"temp_{base_name}.docx"
    doc.save(temp_docx)
    
    # Ensure exam folder exists
    os.makedirs("exam", exist_ok=True)
    
    # Convert to PDF
    temp_pdf = os.path.join("exam", f"temp_{base_name}.pdf")
    convert(temp_docx, temp_pdf)
    
    # Add footer to PDF
    output_pdf = os.path.join("exam", f"{base_name}.pdf")
    add_footer_to_pdf(temp_pdf, output_pdf, base_name)
    
    # Remove temporary files
    os.remove(temp_docx)
    os.remove(temp_pdf)
    
    print(f"Processing complete. PDF saved at: {output_pdf}")

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def add_footer_to_pdf(input_pdf, output_pdf, base_name):
    reader = PdfReader(input_pdf)
    writer = PdfWriter()

    for page_num in range(len(reader.pages)):
        page = reader.pages[page_num]
        packet = BytesIO()
        can = canvas.Canvas(packet, pagesize=letter)
        
        # Add footer
        footer_text = f"{base_name} {page_num + 1}/{len(reader.pages)}"
        can.setFont("SimSun", 10)  # 使用 SimSun 字体
        can.drawCentredString(300, 30, footer_text)
        can.save()

        packet.seek(0)
        new_pdf = PdfReader(packet)
        page.merge_page(new_pdf.pages[0])
        writer.add_page(page)

    with open(output_pdf, "wb") as output_file:
        writer.write(output_file)

# Usage example
input_file = "2016考研英语一真题.docx"  # Replace with your input file name
process_docx(input_file)